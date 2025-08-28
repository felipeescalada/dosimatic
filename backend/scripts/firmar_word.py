#!/usr/bin/env python3
"""
Script para insertar una tabla de firmas en un documento Word (.docx)
Requiere: pip install python-docx pillow
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def agregar_tabla_firmas(doc, firmas):
    """
    Inserta una tabla de firmas en el documento
    
    Args:
        doc: objeto Document
        firmas (list): lista de diccionarios con:
            {
              "titulo": "ELABORÓ",
              "cargo": "Director de Gestión y Control de la Calidad",
              "firma": "firma1.png",
              "fecha": "26/06/2017"
            }
    """
    # Crear tabla con 4 filas x 3 columnas
    table = doc.add_table(rows=4, cols=3)
    try:
        table.style = "Table Grid"
    except KeyError:
        # Si el estilo no existe, usar None (default)
        table.style = None

    # Fila 1 -> Títulos
    for i, f in enumerate(firmas):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        run = p.add_run(f["titulo"])
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fila 2 -> Cargo
    for i, f in enumerate(firmas):
        cell = table.cell(1, i)
        p = cell.paragraphs[0]
        p.add_run(f["cargo"]).font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fila 3 -> Firma (imagen o texto)
    for i, f in enumerate(firmas):
        cell = table.cell(2, i)
        p = cell.paragraphs[0]
        if f["firma"] and os.path.exists(f["firma"]):
            run = p.add_run()
            run.add_picture(f["firma"], width=Inches(1.5), height=Inches(0.8))
        else:
            p.add_run("FIRMA DIGITAL").italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fila 4 -> Fechas
    for i, f in enumerate(firmas):
        cell = table.cell(3, i)
        p = cell.paragraphs[0]
        if i == 0:
            p.add_run(f"Fecha de Emisión: {f['fecha']}")
        elif i == 1:
            p.add_run(f"Fecha de Aplicación: {f['fecha']}")
        else:
            p.add_run(f"Vigencia: {f['fecha']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def firmar_word(input_path, output_path, firmas):
    """
    Inserta tabla de firmas en un documento Word
    """
    try:
        print(f"\n🔍 Iniciando proceso:")
        print(f"   - Documento: {input_path}")
        print(f"   - Salida: {output_path}")
        
        if not os.path.exists(input_path):
            print(f"❌ Error: El archivo de entrada no existe: {input_path}")
            return False
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc = Document(input_path)

        # Obtener o crear el pie de página
        section = doc.sections[0]
        footer = section.footer
        
        # Limpiar el pie de página existente
        for elem in footer.paragraphs + footer.tables:
            p = elem._element
            p.getparent().remove(p)
        
        # Insertar la tabla de firmas en el footer
        from docx.shared import Inches
        table = footer.add_table(rows=4, cols=3, width=Inches(6.0))
        
        # Configurar la tabla
        try:
            table.style = "Table Grid"
        except KeyError:
            table.style = None

        # Llenar la tabla con las firmas
        for i, f in enumerate(firmas):
            # Fila 1: Título
            cell = table.cell(0, i)
            p = cell.paragraphs[0]
            run = p.add_run(f["titulo"])
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Fila 2: Cargo
            cell = table.cell(1, i)
            p = cell.paragraphs[0]
            p.add_run(f["cargo"]).font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Fila 3: Firma (imagen o texto)
            cell = table.cell(2, i)
            p = cell.paragraphs[0]
            if f["firma"] and os.path.exists(f["firma"]):
                try:
                    run = p.add_run()
                    run.add_picture(f["firma"], width=Inches(1.5), height=Inches(0.8))
                except Exception as e:
                    print(f"⚠️  Error insertando imagen de firma: {str(e)}")
                    p.add_run("FIRMA DIGITAL").italic = True
            else:
                p.add_run("FIRMA DIGITAL").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Fila 4: Fechas
            cell = table.cell(3, i)
            p = cell.paragraphs[0]
            run = p.add_run()
            if i == 0:
                run.text = f"Fecha Emisión: {f['fecha']}"
            elif i == 1:
                run.text = f"Fecha Aplicación: {f['fecha']}"
            else:
                run.text = f"Vigencia: {f['fecha']}"
            run.font.size = Pt(7)  # Tamaño de letra más pequeño
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Fila 4: Fechas
            #cell = table.cell(3, i)
            #p = cell.paragraphs[0]
            #if i == 0:
            #    p.add_run(f"Fecha de Emisión: {f['fecha']}")
            #elif i == 1:
            #    p.add_run(f"Fecha de Aplicación: {f['fecha']}")
            #else:
            #    p.add_run(f"Vigencia: {f['fecha']}")
            #p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_path)
        print(f"✅ Documento guardado en: {output_path}")
        return True

    except Exception as e:
        print(f"❌ Error en firmar_word: {str(e)}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return False

def main():
    """Función principal del script"""
    if len(sys.argv) != 5:
        print("Uso: python3 firmar_word.py <input_docx> <signature_image> <output_docx> <signer_name>")
        sys.exit(1)

    input_path = os.path.abspath(sys.argv[1])
    signature_path = os.path.abspath(sys.argv[2]) if sys.argv[2].strip() else ""
    output_path = os.path.abspath(sys.argv[3])
    signer_name = sys.argv[4]

    # Definir la estructura de la tabla con los datos recibidos
    firmas = [
        {
            "titulo": "ELABORÓ",
            "cargo": signer_name,
            "firma": signature_path,
            "fecha": datetime.now().strftime("%d/%m/%Y")
        },
        {
            "titulo": "REVISÓ",
            "cargo": "Pendiente",
            "firma": "",
            "fecha": datetime.now().strftime("%d/%m/%Y")
        },
        {
            "titulo": "APROBÓ",
            "cargo": "Pendiente",
            "firma": "",
            "fecha": datetime.now().strftime("%d/%m/%Y")
        }
    ]

    success = firmar_word(input_path, output_path, firmas)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
