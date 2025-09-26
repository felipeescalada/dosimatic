#!/usr/bin/env python3
"""
Script para firmar documentos Word (.docx) insertando una imagen de firma o solo texto
Requiere: pip install python-docx pillow
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from docx.oxml.shared import qn
from docx.oxml import parse_xml

def firmar_word(input_path, signature_path, output_path, signer_name):
    """
    Firma un documento Word insertando una imagen de firma o solo texto
    
    Args:
        input_path (str): Ruta al documento de entrada
        signature_path (str): Ruta a la imagen de firma (puede ser vacía)
        output_path (str): Ruta donde guardar el documento firmado
        signer_name (str): Nombre del firmante
        
    Returns:
        bool: True si tuvo éxito, False si falló
    """
    try:
        print(f"\n🔍 Iniciando firma de documento:")
        print(f"   - Documento: {input_path}")
        print(f"   - Firma: {signature_path if signature_path else 'Solo texto'}")
        print(f"   - Salida: {output_path}")
        print(f"   - Firmante: {signer_name}")
        
        # Verificar si el archivo de entrada existe
        if not os.path.exists(input_path):
            print(f"❌ Error: El archivo de entrada no existe: {input_path}")
            return False
            
        # Verificar si la imagen de firma existe
        has_signature_image = bool(signature_path and os.path.exists(signature_path))
        if signature_path and not has_signature_image:
            print(f"⚠️  Advertencia: No se encontró la imagen de firma en: {signature_path}")
            print(f"    Se usará solo texto para la firma.")
        
        # Crear directorio de salida si no existe
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Cargar el documento
        doc = Document(input_path)
        
        # Obtener la sección de footer (o crearla si no existe)
        section = doc.sections[0]
        footer = section.footer
        
        # Limpiar el footer existente de manera compatible
        for paragraph in footer.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Agregar párrafo simple en el footer para la firma
        if has_signature_image:
            try:
                # Agregar la imagen de firma en el footer
                signature_paragraph = footer.add_paragraph()
                signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Insertar imagen con tamaño controlado
                run = signature_paragraph.add_run()
                run.add_picture(signature_path, width=Inches(1.5), height=Inches(0.8))
                print("✅ Imagen de firma insertada correctamente 1")

                signature_paragraph = footer.add_paragraph()
                signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Insertar imagen con tamaño controlado
                run = signature_paragraph.add_run()
                run.add_picture(signature_path, width=Inches(1.5), height=Inches(0.8))
                print("✅ Imagen de firma insertada correctamente 2")
                
            except Exception as e:
                print(f"⚠️  Error insertando imagen de firma: {str(e)}")
                has_signature_image = False
        
        # Si no hay imagen o falló la inserción, usar texto
        if not has_signature_image:
            signature_paragraph = footer.add_paragraph()
            signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = signature_paragraph.add_run("FIRMA DIGITAL")
            run.bold = True
            run.font.size = Pt(9)
            print("ℹ️  Usando firma de texto")
        
        # Agregar información del firmante
        info_paragraph = footer.add_paragraph()
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Nombre del firmante
        run = info_paragraph.add_run(f"{signer_name}")
        run.bold = True
        run.font.size = Pt(8)
        
        # Fecha y hora
        info_paragraph.add_run("\n")
        fecha_firma = datetime.now().strftime("%d/%m/%Y %H:%M")
        run = info_paragraph.add_run(f"Firmado el {fecha_firma}")
        run.font.size = Pt(7)
        
        # Guardar el documento firmado
        doc.save(output_path)
        
        # Verificar que se creó el archivo
        if not os.path.exists(output_path):
            print(f"❌ Error: No se pudo guardar el archivo en: {output_path}")
            return False
            
        print(f"✅ Documento firmado guardado en: {output_path}")
        print(f"   Tamaño del archivo: {os.path.getsize(output_path) / 1024:.2f} KB")
        return True
        
    except Exception as e:
        print(f"❌ Error en firmar_word: {str(e)}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return False

def main():
    """Función principal del script"""
    if len(sys.argv) != 5:
        print("Uso: python3 firmar_word.py <input_docx> <signature_image> <output_docx> <signer_name>")
        print("Ejemplo: python3 firmar_word.py documento.docx firma.png documento_firmado.docx 'Juan Pérez'")
        sys.exit(1)
    
    input_path = os.path.abspath(sys.argv[1])
    signature_path = os.path.abspath(sys.argv[2]) if sys.argv[2].strip() else ""
    output_path = os.path.abspath(sys.argv[3])
    signer_name = sys.argv[4]
    
    if not os.path.exists(input_path):
        print(f"Error: El archivo de entrada no existe: {input_path}")
        sys.exit(1)
    
    success = firmar_word(input_path, signature_path, output_path, signer_name)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
