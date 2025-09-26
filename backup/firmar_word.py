#!/usr/bin/env python3
"""
Script para firmar documentos Word (.docx) insertando una imagen de firma o solo texto
Requiere: pip install python-docx pillow
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def firmar_word(input_path, signature_path, output_path, signer_name):
    """
    Firma un documento Word insertando una imagen de firma o solo texto
    
    Args:
        input_path (str): Ruta del documento Word original
        signature_path (str): Ruta de la imagen de firma (puede estar vacía)
        output_path (str): Ruta donde guardar el documento firmado
        signer_name (str): Nombre del firmante
    """
    try:
        # Verificar que el archivo de entrada existe
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Archivo no encontrado: {input_path}")
        
        # Verificar si hay imagen de firma
        has_signature_image = signature_path and signature_path.strip() and os.path.exists(signature_path)
        
        # Abrir el documento
        doc = Document(input_path)
        
        # Agregar un salto de página al final
        doc.add_page_break()
        
        # Agregar título de firma
        firma_title = doc.add_paragraph()
        firma_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = firma_title.add_run("DOCUMENTO FIRMADO DIGITALMENTE")
        run.bold = True
        run.font.size = Inches(0.16)  # 12pt aproximadamente
        
        # Agregar espacio
        doc.add_paragraph()
        
        # Si hay imagen de firma, agregarla
        if has_signature_image:
            # Agregar la imagen de firma centrada
            firma_paragraph = doc.add_paragraph()
            firma_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Insertar imagen con tamaño específico
            run = firma_paragraph.runs[0] if firma_paragraph.runs else firma_paragraph.add_run()
            run.add_picture(signature_path, width=Inches(2.5), height=Inches(1.5))
        else:
            # Agregar firma de texto simple
            firma_paragraph = doc.add_paragraph()
            firma_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = firma_paragraph.add_run("FIRMA DIGITAL")
            run.bold = True
            run.font.size = Inches(0.2)  # 14pt aproximadamente
        
        # Agregar información del firmante
        info_paragraph = doc.add_paragraph()
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nombre del firmante
        run = info_paragraph.add_run(f"Firmado por: {signer_name}")
        run.bold = True
        
        # Salto de línea
        info_paragraph.add_run("\n")
        
        # Fecha y hora
        fecha_firma = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        run = info_paragraph.add_run(f"Fecha de firma: {fecha_firma}")
        
        # Agregar línea separadora
        doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Guardar el documento firmado
        doc.save(output_path)
        
        signature_type = "con imagen" if has_signature_image else "solo texto"
        print(f"Documento firmado exitosamente ({signature_type}): {output_path}")
        return True
        
    except Exception as e:
        print(f"Error firmando documento: {str(e)}", file=sys.stderr)
        return False

def main():
    """Función principal del script"""
    if len(sys.argv) != 5:
        print("Uso: python3 firmar_word.py <input_docx> <signature_image> <output_docx> <signer_name>")
        print("Ejemplo: python3 firmar_word.py documento.docx firma.png documento_firmado.docx 'Juan Pérez'")
        sys.exit(1)
    
    input_path = sys.argv[1]
    signature_path = sys.argv[2]
    output_path = sys.argv[3]
    signer_name = sys.argv[4]
    
    # Crear directorio de salida si no existe
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    success = firmar_word(input_path, signature_path, output_path, signer_name)
    
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main()
